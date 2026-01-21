from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Iterable, Optional

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas


@dataclass
class GeneratedFile:
    filename: str
    content_type: str
    data: bytes


class DocumentService:
    """
    Generates downloadable documents (DOCX/PDF) for:
    - Cover letter
    - Enhanced CV (text-based version)
    """

    # -----------------------
    # Cover letter
    # -----------------------
    def cover_letter_docx(self, cover_letter_text: str, filename: str = "cover_letter.docx") -> GeneratedFile:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for para in cover_letter_text.strip().split("\n"):
            doc.add_paragraph(para)

        buf = BytesIO()
        doc.save(buf)
        return GeneratedFile(filename=filename, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", data=buf.getvalue())

    def cover_letter_pdf(self, cover_letter_text: str, filename: str = "cover_letter.pdf") -> GeneratedFile:
        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=LETTER)
        width, height = LETTER

        textobject = c.beginText(54, height - 72)  # margins
        textobject.setFont("Times-Roman", 11)

        for line in cover_letter_text.strip().split("\n"):
            # basic wrapping (simple + robust)
            for wrapped in self._wrap_line(line, max_chars=95):
                textobject.textLine(wrapped)
        c.drawText(textobject)
        c.showPage()
        c.save()

        return GeneratedFile(filename=filename, content_type="application/pdf", data=buf.getvalue())

    # -----------------------
    # Enhanced CV
    # -----------------------
    def enhanced_cv_docx(
        self,
        original_cv_text: str,
        suggestions: Iterable[str],
        filename: str = "enhanced_cv.docx",
        title: Optional[str] = "Enhanced CV (Suggestions)",
    ) -> GeneratedFile:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        if title:
            doc.add_heading(title, level=1)

        doc.add_heading("Original CV (Text Extract)", level=2)
        for para in original_cv_text.strip().split("\n"):
            doc.add_paragraph(para)

        doc.add_heading("Suggested Improvements", level=2)
        for s in suggestions:
            doc.add_paragraph(s, style="List Bullet")

        buf = BytesIO()
        doc.save(buf)
        return GeneratedFile(filename=filename, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", data=buf.getvalue())

    def enhanced_cv_pdf(
        self,
        original_cv_text: str,
        suggestions: Iterable[str],
        filename: str = "enhanced_cv.pdf",
    ) -> GeneratedFile:
        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=LETTER)
        width, height = LETTER

        textobject = c.beginText(54, height - 72)
        textobject.setFont("Times-Roman", 11)

        textobject.textLine("Enhanced CV (Suggestions)")
        textobject.textLine("")
        textobject.textLine("Original CV (Text Extract)")
        textobject.textLine("-" * 30)

        for line in original_cv_text.strip().split("\n"):
            for wrapped in self._wrap_line(line, 95):
                textobject.textLine(wrapped)

        textobject.textLine("")
        textobject.textLine("Suggested Improvements")
        textobject.textLine("-" * 30)

        for s in suggestions:
            for wrapped in self._wrap_line(f"• {s}", 95):
                textobject.textLine(wrapped)

        c.drawText(textobject)
        c.showPage()
        c.save()

        return GeneratedFile(filename=filename, content_type="application/pdf", data=buf.getvalue())

    # -----------------------
    # helpers
    # -----------------------
    def _wrap_line(self, line: str, max_chars: int) -> list[str]:
        if len(line) <= max_chars:
            return [line]
        out: list[str] = []
        words = line.split(" ")
        cur = ""
        for w in words:
            if len(cur) + len(w) + 1 <= max_chars:
                cur = (cur + " " + w).strip()
            else:
                if cur:
                    out.append(cur)
                cur = w
        if cur:
            out.append(cur)
        return out
